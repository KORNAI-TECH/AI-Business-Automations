# Version: 6.7 - Exact Spacing & Maximum Verbatim Length
import os
import asyncio
import logging
import multiprocessing
import re
from datetime import datetime
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.types import FSInputFile, InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.fsm.context import FSMContext
from aiogram.fsm.storage.memory import MemoryStorage
from groq import Groq
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt

# 1. Настройки
load_dotenv()
TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

groq_client = Groq(api_key=GROQ_API_KEY, timeout=400.0)
genai.configure(api_key=GEMINI_API_KEY)
gemini_model = genai.GenerativeModel('gemini-1.5-flash')

bot = Bot(token=TELEGRAM_TOKEN)
dp = Dispatcher(storage=MemoryStorage())
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')

# --- УНИВЕРСАЛЬНЫЙ ПРОМПТ СТРУКТУРЫ ---

def get_meeting_prompt(current_date):
    return f"""
    Твоя задача — составить подробный профессиональный протокол.
    ИСПОЛЬЗУЙ МАРКЕР [SPLIT] ДЛЯ РАЗДЕЛЕНИЯ ТЕКСТА ДЛЯ ЧАТА И ФАЙЛА.

    🗓 Дата: {current_date}

    📂 ТЕМА ВСТРЕЧИ
    [Информативное название]

    👥 УЧАСТНИКИ
    [Список вертикально: — Имя: Роль]

    💡 ВАЖНЫЕ МОМЕНТЫ
    [Ключевые тезисы списком]

    ✅ СПИСОК ЗАДАЧ
    (Оформи блоками по 3 строки. Каждое поле задачи на новой строке.
    Между блоками задач — ОБЯЗАТЕЛЬНО ОДНА пустая строка. Нумерация с 1.)
    1. Задача: ...
       Ответственный: ...
       Сроки: ...

    [SPLIT]
    📝 ДЕТАЛЬНОЕ РЕЗЮМЕ ВСТРЕЧИ
    [Связный глубокий текст в 3-4 абзаца о ходе встречи и итогах.]

    __________________________________________________________________
    🎤 СТРУКТУРИРОВАННАЯ СТЕНОГРАММА (ПО РОЛЯМ)
    ВАЖНО: Твоя цель — ПОЛНАЯ ДОСЛОВНАЯ расшифровка.
    1. КАТЕГОРИЧЕСКИ ЗАПРЕЩЕНО сокращать или выкидывать реплики.
    2. Перепиши КАЖДОЕ предложение из оригинала. Формат 'Имя: Реплика'.
    3. Каждая новая фраза - с новой строки. Не объединяй реплики!
    4. Имена определяй по контексту обращений. НАЧИНАЙ С ПЕРВОГО СЛОВА.
    """

# --- ЛОГИКА ОБРАБОТКИ ---

async def transcribe_audio(file_path: str):
    try:
        with open(file_path, "rb") as file:
            return groq_client.audio.transcriptions.create(
                file=(file_path, file.read()),
                model="whisper-large-v3",
                response_format="text",
                language="ru",
                prompt="Алло. Здравствуйте. Да, слушаю вас. Соединяю. Начало разговора."
            )
    except Exception as e:
        logging.error(f"Whisper Error: {e}")
        return ""

async def analyze_and_diarize(raw_text: str):
    current_date = datetime.now().strftime("%d %B %Y г.")
    prompt = get_meeting_prompt(current_date) + f"\n\nТЕКСТ ДЛЯ АНАЛИЗА:\n{raw_text}"

    # Сначала Gemini (она лучше держит объем)
    try:
        logging.info("Анализ через Gemini...")
        response = await asyncio.to_thread(gemini_model.generate_content, prompt)
        return response.text.replace("**", "").replace("__", ""), "Gemini"
    except Exception:
        # Fallback на Groq
        response = groq_client.chat.completions.create(
            messages=[{"role": "system", "content": "Ты элитный стенографист. Пиши дословно. Используй ровно одну пустую строку между разделами."},
                      {"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.0,
            max_tokens=8000 # Лимит на максимум
        )
        return response.choices[0].message.content.replace("**", "").replace("__", ""), "Groq Llama"

# --- ОБРАБОТКА ТЕКСТА И ФАЙЛА ---

def normalize_spacing(text):
    """Удаляет лишние пустые строки, оставляя ровно одну"""
    # Заменяем 3 и более переносов на 2 (что дает одну пустую строку)
    return re.sub(r'\n{3,}', '\n\n', text.strip())

def add_styled_para(doc, line):
    line = line.strip()
    if not line:
        doc.add_paragraph()
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    if any(emoji in line for emoji in ["🗓", "📂", "👥", "📝", "💡", "✅", "🎤"]) or "Дата:" in line:
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(12)
    elif ":" in line and len(line.split(":")[0]) < 45:
        parts = line.split(":", 1)
        run_name = p.add_run(f"{parts[0].strip()}: ")
        run_name.bold = True
        if len(parts) > 1:
            p.add_run(parts[1])
    else:
        p.add_run(line)

def create_docx(report_text, filename):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    parts = report_text.split("[SPLIT]")
    # 1. Чат
    for line in parts[0].strip().split('\n'):
        add_styled_para(doc, line)
    # 2. Файл (Резюме и Стенограмма)
    if len(parts) > 1:
        doc.add_page_break()
        second_part = parts[1].strip()
        for line in second_part.split('\n'):
            add_styled_para(doc, line)
    doc.save(filename)

# --- ХЕНДЛЕРЫ ---

@dp.message(Command("start"))
async def cmd_start(message: types.Message):
    await message.answer("🚀 Бот v6.7 готов. Исправлены отступы и увеличена длина стенограммы. Жду аудио!")

@dp.message(F.voice | F.audio | F.video_note)
async def handle_audio(message: types.Message, state: FSMContext):
    status_msg = await message.answer("⏳ Анализирую аудио...")
    ts = datetime.now().strftime("%H%M%S")
    audio_path = f"temp_{ts}.ogg"

    try:
        obj = message.voice or message.audio or message.video_note
        await bot.download_file((await bot.get_file(obj.file_id)).file_path, audio_path)

        raw_text = await transcribe_audio(audio_path)
        if not raw_text:
            await status_msg.edit_text("❌ Ошибка расшифровки.")
            return

        final_report, engine = await analyze_and_diarize(raw_text)

        # ЧИСТКА ОТСТУПОВ ДЛЯ ТЕЛЕГРАМ
        full_cleaned = normalize_spacing(final_report)
        chat_text = full_cleaned.split("[SPLIT]")[0].strip()

        await state.update_data(report=full_cleaned)

        kb = InlineKeyboardMarkup(inline_keyboard=[[InlineKeyboardButton(text="📝 Скачать полную версию (Word)", callback_data="get_word")]])

        await message.answer(f"📊 <b>ИТОГИ ({engine}):</b>\n\n{chat_text}", parse_mode="HTML", reply_markup=kb)
        await status_msg.delete()
    except Exception as e:
        await message.answer(f"⚠️ Ошибка: {e}")
    finally:
        if os.path.exists(audio_path):
            os.remove(audio_path)

@dp.callback_query(F.data == "get_word")
async def send_file(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()
    report = data.get("report")
    if not report:
        await callback.answer("Данные не найдены.", show_alert=True)
        return

    filename = f"Protocol_{datetime.now().strftime('%H%M%S')}.docx"
    try:
        create_docx(report, filename)
        await callback.message.answer_document(FSInputFile(filename), caption="📄 Полный профессиональный протокол")
        await callback.answer()
    finally:
        await asyncio.sleep(10)
        if os.path.exists(filename):
            os.remove(filename)

async def main():
    await dp.start_polling(bot, skip_updates=True)

def run_dummy_server():
    from http.server import HTTPServer, BaseHTTPRequestHandler
    class Handler(BaseHTTPRequestHandler):
        def do_GET(self):
            self.send_response(200)
            self.end_headers()
            self.wfile.write(b"OK")

    port = int(os.environ.get("PORT", 8080))
    HTTPServer(('0.0.0.0', port), Handler).serve_forever()

if __name__ == "__main__":
    multiprocessing.Process(target=run_dummy_server, daemon=True).start()
    asyncio.run(main())